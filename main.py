from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from typing import Optional
from PIL import Image
import pytesseract
import pdfplumber
import io
from docx import Document as DocxDocument
from langchain.docstore.document import Document
from langchain_community.chat_models import ChatOllama
from langchain.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.messages import HumanMessage
from fastapi.middleware.cors import CORSMiddleware
import pypandoc

app = FastAPI()

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Load the language model
# local_llm = 'gemma'
local_llm = 'llama3'
# local_llm = 'llama3.1'
# local_llm = 'mistral'

# --------linus
# host_ip = os.environ.get('HOST_IP', 'host.docker.internal')
# llm = ChatOllama(model=local_llm, temperature=0, base_url=f"http://{host_ip}:11434")
# llm = ChatOllama(model=local_llm, temperature=0)

# --------mac
# Load the language model
llm = ChatOllama(model=local_llm, temperature=0, base_url="http://host.docker.internal:11434")

# Define the QA prompt template
qa_system_prompt = """You're a friendly expert in math and finance. Given the context, answer the question, and provide a clear, brief explanation afterward.
Question: {input}
Context: {context}
Answer: assistant"""
qa_prompt = ChatPromptTemplate.from_messages(
    [
        ("system", qa_system_prompt),
        MessagesPlaceholder("chat_history"),
        ("human", "{input}"),
    ]
)

# Create the question-answer chain
question_answer_chain = create_stuff_documents_chain(llm, qa_prompt)

# Chat history
chat_history = []

class QueryModel(BaseModel):
    question: str
    image: Optional[UploadFile] = None

# Function to extract text from an image
def extract_text_from_image(image):
    return pytesseract.image_to_string(image)

# Function to extract text from a PDF
def extract_text_from_pdf(pdf_bytes):
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        extracted_text = ""
        for page in pdf.pages:
            extracted_text += page.extract_text()
    return extracted_text

# Function to extract text from a DOCX file
def extract_text_from_docx(docx_bytes):
    doc = DocxDocument(io.BytesIO(docx_bytes))
    extracted_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return extracted_text

# Function to extract text from a DOC file using pypandoc
def extract_text_from_doc(doc_bytes):
    with io.BytesIO(doc_bytes) as doc_file:
        extracted_text = pypandoc.convert_file(doc_file.name, 'plain', format='doc')
    return extracted_text

@app.post("/analyze-file/")
async def analyze_file(question: str = Form(...), file: UploadFile = File(...)):
    try:
        # Determine the file type by its content type
        file_content_type = file.content_type

        if file_content_type == "application/pdf":
            # Extract text from PDF
            pdf_bytes = await file.read()
            extracted_text = extract_text_from_pdf(pdf_bytes)
        elif "image" in file_content_type:
            # Extract text from Image
            image = Image.open(io.BytesIO(await file.read()))
            extracted_text = extract_text_from_image(image)
        elif file_content_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
            # Extract text from DOCX or DOC
            doc_bytes = await file.read()
            if file_content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                extracted_text = extract_text_from_docx(doc_bytes)
            elif file_content_type == "application/msword":
                extracted_text = extract_text_from_doc(doc_bytes)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Please upload a PDF, image, or Word file.")

        # Create a context document from the extracted text
        context_document = Document(page_content=extracted_text)

        # Ask the question
        ai_msg_1 = question_answer_chain.invoke({"input": question, "chat_history": [], "context": [context_document]})
        chat_history.extend([HumanMessage(content=question), ai_msg_1])

        return JSONResponse(content={"answer": ai_msg_1})
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=3000)
